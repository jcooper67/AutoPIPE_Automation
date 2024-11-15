import tkinter as tk
from tkinter import filedialog
import math 

import pandas as pd
from docx import Document

user_support_nodes = []
user_nozzle_nodes = []
user_transform = ""
swap_order = []
swap_scalar = []

doc = Document("C:\\Users\\COOPERJ\\Documents\\Projects\\AutoPIPE_Automation\\Sample Files\\Table_Template1.docx")

def transform_nozzle_loads(nozzle_forces, nozzle_moments, swap_order, swap_scalar):
    new_forces = []
    new_moments = []
    for i in range(len(nozzle_forces)):
        new_force = (nozzle_forces[swap_order[i]])*swap_scalar[i]
        new_forces.append(new_force)
        new_moment = (nozzle_moments[swap_order[i]])*swap_scalar[i]
        new_moments.append(new_moment)
    return new_forces, new_moments
    

def process_transfromation():
    global user_transform
    swap_order,swap_scalar = [],[]
    user_transform = user_transform.upper()


    axis_key = {
        "X":0,
        "Y":1,
        "Z":2
    }

    i = 0
    while i<(len(user_transform)):
        if user_transform[i] == "-":
            swap_scalar.append(-1)
            swap_order.append(axis_key[user_transform[i+1]])
            i+=2
        else:
            swap_scalar.append(1)
            swap_order.append(axis_key[user_transform[i]])
            i+=1
    return swap_order,swap_scalar


def update_nozzle_tables(nozzle_dict):
    for i, key in enumerate(nozzle_dict):
        nozzle_list = nozzle_dict[key]
        table = doc.tables[i]
        
        second_row = table.rows[1]

        second_row.cells[9].text = key
        twelfth_row = table.rows[11]

        #Deadweight Forces
        twelfth_row.cells[2].text = str(nozzle_list[0][0])
        twelfth_row.cells[3].text = str(nozzle_list[0][1])
        twelfth_row.cells[4].text = str(nozzle_list[0][2])

        #Calc Deadweight FSR
        deadweight_fsr = math.sqrt(((nozzle_list[0][1])**2)+((nozzle_list[0][2])**2))
        deadweight_fsr = round(deadweight_fsr,0)

        twelfth_row.cells[5].text = str(deadweight_fsr)


        #Deadweight Moments
        twelfth_row.cells[6].text = str(nozzle_list[0][3])
        twelfth_row.cells[7].text = str(nozzle_list[0][4])
        twelfth_row.cells[8].text = str(nozzle_list[0][5])

        deadweight_msr = math.sqrt(((nozzle_list[0][4])**2)+((nozzle_list[0][5])**2))
        deadweight_msr = round(deadweight_msr,0)

        twelfth_row.cells[9].text = str(deadweight_msr)

        #Th-1 Forces
        thirteenth_row = table.rows[12]
        thirteenth_row.cells[2].text = str(nozzle_list[1][0])
        thirteenth_row.cells[3].text = str(nozzle_list[1][1])
        thirteenth_row.cells[4].text = str(nozzle_list[1][2])

        th1_fsr = math.sqrt((nozzle_list[1][1])**2+(nozzle_list[1][2])**2)
        th1_fsr = round(th1_fsr,0)

        thirteenth_row.cells[5].text = str(th1_fsr)



        #Th-1 Moments
        thirteenth_row.cells[6].text = str(nozzle_list[1][3])
        thirteenth_row.cells[7].text = str(nozzle_list[1][4])
        thirteenth_row.cells[8].text = str(nozzle_list[1][5])

        th1_msr = math.sqrt((nozzle_list[1][4])**2+(nozzle_list[1][5])**2)
        th1_msr = round(th1_msr,0)

        thirteenth_row.cells[9].text = str(th1_msr)




        #Th-2 Forces
        fourteenth_row = table.rows[13]
        fourteenth_row.cells[2].text = str(nozzle_list[2][0])
        fourteenth_row.cells[3].text = str(nozzle_list[2][1])
        fourteenth_row.cells[4].text = str(nozzle_list[2][2])

        th2_fsr = math.sqrt((nozzle_list[2][1])**2+(nozzle_list[2][2])**2)
        th2_fsr = round(th2_fsr,0)

        fourteenth_row.cells[5].text = str(th2_fsr)

        
        #Th-2 Moments
        fourteenth_row.cells[6].text = str(nozzle_list[2][3])
        fourteenth_row.cells[7].text = str(nozzle_list[2][4])
        fourteenth_row.cells[8].text = str(nozzle_list[2][5])

        th2_msr = math.sqrt((nozzle_list[2][4])**2+(nozzle_list[2][5])**2)
        th2_msr = round(th2_msr,0)

        fourteenth_row.cells[9].text = str(th2_msr)

        #Calculate Hot(Weight + Envelope of Expansion Cases)

        fifthteenth_row = table.rows[14]

        deadweight_list = nozzle_list[0]

        thermal1_list = nozzle_list[1]
        thermal2_list = nozzle_list[2]

        envelope_thermal = [thermal1_list[i] if abs(thermal1_list[i]) >= abs(thermal2_list[i]) else thermal2_list[i] for i in range(len(thermal1_list))]
        
        hot_list = [deadweight_list[i]+envelope_thermal[i] for i in range(len(deadweight_list))]

        hot_fsr = deadweight_fsr + max(abs(th1_fsr),abs(th2_fsr))

        hot_msr = deadweight_msr + max(abs(th1_msr),abs(th2_msr))


        fifthteenth_row.cells[2].text = str(hot_list[0])
        fifthteenth_row.cells[3].text = str(hot_list[1])
        fifthteenth_row.cells[4].text = str(hot_list[2])
        fifthteenth_row.cells[5].text = str(hot_fsr)
        
        #Th-2 Moments
        fifthteenth_row.cells[6].text = str(hot_list[3])
        fifthteenth_row.cells[7].text = str(hot_list[4])
        fifthteenth_row.cells[8].text = str(hot_list[5])
        fifthteenth_row.cells[9].text = str(hot_msr)


    #Calc Maximum of DW,Hot

        max_list = [max(abs(hot_list), abs(deadweight_list)) for hot_list, deadweight_list in zip(hot_list, deadweight_list)]

        seventeenth_row = table.rows[16]

        seventeenth_row.cells[2].text = str(max_list[0])

        max_fsr = max(abs(deadweight_fsr),abs(hot_fsr))

        seventeenth_row.cells[5].text = str(max_fsr)

        seventeenth_row.cells[6].text = str(max_list[3])

        max_msr = max(abs(deadweight_msr),abs(hot_msr))

        seventeenth_row.cells[9].text = str(max_msr)
        






def update_support_table(support_dict):
    table = doc.tables[3]
    for i, key in enumerate(support_dict):
        
        support_list = support_dict[key]
        #print (support_list)
        first_row = table.rows[i*3+1]
        second_row = table.rows[i*3+2]
        third_row = table.rows[i*3+3]
        first_row.cells[0].text = key
        first_row.cells[1].text = str(support_list[0])
        second_row.cells[3].text = str(support_list[1])
        second_row.cells[4].text = str(support_list[2])


def process_data():

    support_dict = {}

    df = pd.read_excel(file_path, sheet_name='Support_Forces')
    df.columns = df.iloc[0]
    df=df[1:]
    df = df.reset_index(drop = True)

    for support_node in user_support_nodes:

        support_dict[support_node] = []

        filtered_df = df[(df.iloc[:,0] == int(support_node))]
        support_number = filtered_df.iloc[0]["Tag_No"]

        support_dict[support_node].append(support_number)


        
        filtered_df = df[(df.iloc[:,0] == int(support_node)) & (df["Load_Combination"]== 'GRT{1}')]

        if not filtered_df.empty:
             # Find the row with the max absolute value in the 12th column (Global_Force)
            max_index = filtered_df['Global_Force'].abs().idxmax()
            # Now get the value from the original DataFrame using the max index
            global_force_value = df.loc[max_index,'Global_Force']
            global_force_value = round(global_force_value,0)
            #print(global_force_value)
            support_dict[support_node].append(global_force_value)

        else:
            print("No data found for the specified conditions.")

        filtered_df = df[(df.iloc[:,0] == int(support_node)) & (df["Load_Combination"]== 'GRT{2}')]

        if not filtered_df.empty:
        # Find the row with the max absolute value in the 12th column (Global_Force)
            max_index = filtered_df['Global_Force'].abs().idxmax()
            # Now get the value from the original DataFrame using the max index
            global_force_value = df.loc[max_index,'Global_Force']
            global_force_value= round(global_force_value,0)
            support_dict[support_node].append(global_force_value)
        
        else:
            print("No data found for the specified conditions.")

        
        
        if abs(support_dict[support_node][1])<abs(support_dict[support_node][2]):
            temp_value = support_dict[support_node][2]
            support_dict[support_node][2] = support_dict[support_node][1]
            support_dict[support_node][1] = temp_value

        print(support_dict)

    df = pd.read_excel(file_path, sheet_name='Restraint_Loads')
    df.columns = df.iloc[0]
    df = df[1:]
    df = df.reset_index(drop = True)
    nozzle_dict = {}

    for nozzle_node in user_nozzle_nodes:
        nozzle_dict[nozzle_node] = []
        

        filtered_df = df[(df.iloc[:,0] == int(nozzle_node)) & (df["Load_Combination"]== 'Gravity{1}')]
        forces = filtered_df.iloc[0][["Forces_X","Forces_Y","Forces_Z"]].tolist()
        moments = filtered_df.iloc[0][["Moments_X","Moments_Y","Moments_Z"]].tolist()

       

        forces = [round(x,0) for x in forces]
        moments = [round(x,0) for x in moments]


       

        

        if user_transform:
            swap_order,swap_scalar = process_transfromation()
           
            forces,moments = transform_nozzle_loads(forces,moments,swap_order,swap_scalar)
            

        combined_values = forces+moments
        
        # Append the combined values to the dictionary under the nozzle node
        nozzle_dict[nozzle_node].append(combined_values)

        filtered_df = df[(df.iloc[:,0] == int(nozzle_node)) & (df["Load_Combination"]== 'Thermal 1{1}')]

        forces = filtered_df.iloc[0][["Forces_X","Forces_Y","Forces_Z"]].tolist()
        moments = filtered_df.iloc[0][["Moments_X","Moments_Y","Moments_Z"]].tolist()
        forces = [round(x,0) for x in forces]
        moments = [round(x,0) for x in moments]

        if user_transform:
            swap_order,swap_scalar = process_transfromation()
            forces,moments = transform_nozzle_loads(forces,moments,swap_order,swap_scalar)



        combined_values = forces+moments
        
        # Append the combined values to the dictionary under the nozzle node
        nozzle_dict[nozzle_node].append(combined_values)

        filtered_df = df[(df.iloc[:,0] == int(nozzle_node)) & (df["Load_Combination"]== 'Thermal 2{1}')]
        forces = filtered_df.iloc[0][["Forces_X","Forces_Y","Forces_Z"]].tolist()
        moments = filtered_df.iloc[0][["Moments_X","Moments_Y","Moments_Z"]].tolist()

        forces = [round(x,0) for x in forces]
        moments = [round(x,0) for x in moments]

        if user_transform:
            swap_order,swap_scalar = process_transfromation()
            forces,moments = transform_nozzle_loads(forces,moments,swap_order,swap_scalar)



        combined_values = forces+moments
        
        # Append the combined values to the dictionary under the nozzle node
        nozzle_dict[nozzle_node].append(combined_values)
      

        print(nozzle_dict)



    update_support_table(support_dict)
    update_nozzle_tables(nozzle_dict)
   
    doc.save("C:\\Users\\COOPERJ\\Documents\\Projects\\AutoPIPE_Automation\\Sample Files\\Table_Template1.docx")






def on_submit():

    global user_support_nodes, user_nozzle_nodes, user_transform
    # Get the values from the textboxes
    user_nozzle_nodes = [item.strip() for item in user_nozzles.get().replace(',', ' ').split() if item.strip()]
    user_support_nodes = [item.strip() for item in user_supports.get().replace(',', ' ').split() if item.strip()] 
    user_transform = user_transform.get()      
    print(f"Nozzle Nodes: {user_nozzle_nodes}")
    print(f"Support Nodes: {user_support_nodes}")
    process_data()
    input_window.destroy()  # Close the input window after submission

root = tk.Tk()
root.withdraw()  # Hide the root window

# Ask for the file path
file_path = filedialog.askopenfilename(
    title="Select a file",  # Title of the dialog
    filetypes=[ ("Excel Files", "*.xlsx"),  # Filter for .xlsx files
              ("Excel Files", "*.xls"),  # Filter for .xls files
              ("All Files", "*.*")]  # Show all file types as a fallback
)

# Check if a file was selected
if file_path:
    print(f"File selected: {file_path}")
    
    # Create a new window for user input
    input_window = tk.Toplevel(root)
    input_window.title("Node Number Input")

    # Add two textboxes (Entry widgets)
    nozzle_input = tk.Label(input_window, text="Nozzle Node Numbers (Comma or Space Separated):")
    nozzle_input.pack(padx=10, pady=5)
    user_nozzles = tk.Entry(input_window)
    user_nozzles.pack(padx=10, pady=5)

    support_input = tk.Label(input_window, text="Support Node Numbers (Comma or Space Separated):")
    support_input.pack(padx=10, pady=5)
    user_supports = tk.Entry(input_window)
    user_supports.pack(padx=10, pady=5)

    transform_input = tk.Label(input_window,text = "Coordinate Transformation Assuming XYZ")
    transform_input.pack(padx=10,pady=5)
    user_transform = tk.Entry(input_window)
    user_transform.pack(padx=10,pady=5)

    # Add a submit button to capture the input
    submit_button = tk.Button(input_window, text="Submit", command=on_submit)
    submit_button.pack(padx=10, pady=10)

    # Run the new input window
    input_window.mainloop()

else:
    print("No file selected")



