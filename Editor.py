from docx import Document
import math
from tkinter import messagebox
import tkinter as tk
import sys
class Editor:
    def __init__(self,output_path):
        self.doc = Document(output_path)
        self.output_path = output_path
        #Assumes the support table is the 16th table
        self.support_table_location = 15

    def update_support_table(self, support_dict):
        try:
            table = self.doc.tables[self.support_table_location]
            for i, key in enumerate(support_dict):
                support_list = support_dict[key]
                first_row = table.rows[i*3+1]
                first_row.cells[0].text = key
                first_row.cells[1].text = str(support_list[0])

                for j in range(1,4):
                    if support_list[j][0] == 0:
                        continue
                    else:
                        row = table.rows[i*3+j]
                        row.cells[3].text = str(support_list[j][0])
                        row.cells[4].text = str(support_list[j][1])

        except Exception as e:
            print(f"Error: {e}")
            # Show the error message in a Tkinter popup
            root = tk.Tk()
            root.withdraw()  # Hide the root window
            messagebox.showerror("Error", str(e))
            sys.exit()  # Exit the program
            return None
            


    def update_nozzle_tables(self, nozzle_dict):
        try:
            for i, key in enumerate(nozzle_dict):

                nozzle_list = nozzle_dict[key]
                table = self.doc.tables[i]
                
                second_row = table.rows[1]
                second_row.cells[9].text = key
                
                twelfth_row = table.rows[11]
                #Deadweight Forces
                twelfth_row.cells[2].text = str(nozzle_list[0][0])
                twelfth_row.cells[3].text = str(nozzle_list[0][1])
                twelfth_row.cells[4].text = str(nozzle_list[0][2])

                #Calc Deadweight FSR
                deadweight_fsr = math.sqrt(((nozzle_list[0][1])**2)+((nozzle_list[0][2])**2))
                deadweight_fsr = int(round(deadweight_fsr,0))
                twelfth_row.cells[5].text = str(deadweight_fsr)


                #Deadweight Moments
                twelfth_row.cells[6].text = str(nozzle_list[0][3])
                twelfth_row.cells[7].text = str(nozzle_list[0][4])
                twelfth_row.cells[8].text = str(nozzle_list[0][5])

                deadweight_msr = math.sqrt(((nozzle_list[0][4])**2)+((nozzle_list[0][5])**2))
                deadweight_msr = int(round(deadweight_msr,0))

                twelfth_row.cells[9].text = str(deadweight_msr)


                for i in range(self.number_thermal_cases):
                    row = table.rows[12+i]

                    #Forces
                    row.cells[2].text = str(nozzle_list[i+1][0])
                    row.cells[3].text = str(nozzle_list[i+1][1])
                    row.cells[4].text = str(nozzle_list[i+1][2])

                    th1_fsr = math.sqrt((nozzle_list[i+1][1])**2+(nozzle_list[i+1][2])**2)
                    th1_fsr = int(round(th1_fsr,0))

                    row.cells[5].text = str(th1_fsr)


                    #Moments

                    row.cells[6].text = str(nozzle_list[i+1][3])
                    row.cells[7].text = str(nozzle_list[i+1][4])
                    row.cells[8].text = str(nozzle_list[i+1][5])

                    th1_msr = math.sqrt((nozzle_list[i+1][4])**2+(nozzle_list[i+1][5])**2)
                    th1_msr = int(round(th1_msr,0))
                    row.cells[9].text = str(th1_msr)


                #Calculate Hot(Weight + Envelope of Expansion Cases)

                eighteenth_row = table.rows[17]

                deadweight_list = nozzle_list[0]

                thermal1_list = nozzle_list[1]
                thermal2_list = nozzle_list[2]


                thermal_list = []
                for i in range(self.number_thermal_cases):
                    thermal_list.append(nozzle_list[i+1])

                hot_list = [max(values, key=abs) for values in zip(*thermal_list)]
                print(hot_list)

                #hot_list = [a+b for a,b in zip(deadweight_list,hot_list)]

                hot_list = [a + b if abs(a + b) >= abs(a) else a for a, b in zip(deadweight_list, hot_list)]
                


                ## FSR and MSR to be calculated as resultant of hot x and hot z
                hot_fsr = math.sqrt((hot_list[1])**2+(hot_list[2])**2)
                hot_fsr = int(round(hot_fsr,0))
                hot_msr = math.sqrt((hot_list[4])**2+(hot_list[5])**2)
                hot_msr = int(round(hot_msr,0))


                eighteenth_row.cells[2].text = str(hot_list[0])
                eighteenth_row.cells[3].text = str(hot_list[1])
                eighteenth_row.cells[4].text = str(hot_list[2])
                eighteenth_row.cells[5].text = str(hot_fsr)
                
                #Th-2 Moments
                eighteenth_row.cells[6].text = str(hot_list[3])
                eighteenth_row.cells[7].text = str(hot_list[4])
                eighteenth_row.cells[8].text = str(hot_list[5])
                eighteenth_row.cells[9].text = str(hot_msr)


            #Calc Maximum of DW,Hot

                max_list = [max(abs(hot_list), abs(deadweight_list)) for hot_list, deadweight_list in zip(hot_list, deadweight_list)]

                twentieth_row = table.rows[19]

                twentieth_row.cells[2].text = str(max_list[0])

                max_fsr = max(abs(deadweight_fsr),abs(hot_fsr))

                twentieth_row.cells[5].text = str(max_fsr)

                twentieth_row.cells[6].text = str(max_list[3])

                max_msr = max(abs(deadweight_msr),abs(hot_msr))

                twentieth_row.cells[9].text = str(max_msr)
        except Exception as e:
            print(f"Error: {e}")
            # Show the error message in a Tkinter popup
            root = tk.Tk()
            root.withdraw()  # Hide the root window
            messagebox.showerror("Error", str(e))
            sys.exit()  # Exit the program
            return None
        
    def set_number_thermal_cases(self,num_cases):
        self.number_thermal_cases = num_cases
            

    def save_document(self):
        self.doc.save(self.output_path)