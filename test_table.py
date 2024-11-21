from docx import Document

doc = Document("C:\\Users\\COOPERJ\\Documents\\Projects\\AutoPIPE_Automation\\Table_Duplicate.docx")


# Load the document


# Find the first table in the document
tables = doc.tables

# Check if the document contains tables
if tables:
    # Select the first table to duplicate
    table_to_duplicate = tables[0]

    # Create a new table by copying the rows and cells
    new_table = doc.add_table(rows=0, cols=len(table_to_duplicate.columns))

    # Loop through each row of the original table
    for row in table_to_duplicate.rows:
        # Add a row to the new table
        new_row = new_table.add_row()

        # Loop through each cell in the row and copy the content
        for i, cell in enumerate(row.cells):
            new_row.cells[i].text = cell.text

    # Save the modified document
    doc.save('modified_document.docx')
else:
    print("No tables found in the document.")
